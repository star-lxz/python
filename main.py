import tkinter as tk
from tkinter import ttk
from docx import Document
from openpyxl import load_workbook
import os
import tkinter.messagebox as messagebox
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

class DynamicProcedureApp:
    def __init__(self, root):
        self.root = root
        self.root.title("动态规程")

        # 添加下拉选项，输入需要检定的仪器名称
        self.instrument_label = ttk.Label(root, text="请输入需要检定的仪器名称")
        self.instrument_label.grid(row=0, column=0, padx=10, pady=5)
        self.instrument_combobox = ttk.Combobox(root, values=["光电轴角编码器", "角度仪"])
        self.instrument_combobox.grid(row=0, column=1, padx=10, pady=5)
        self.instrument_combobox.bind("<<ComboboxSelected>>", self.create_word_document)

        # 添加等级选择下拉菜单
        self.grade_label = ttk.Label(root, text="请选择等级")
        self.grade_label.grid(row=0, column=2, padx=10, pady=5)
        self.grade_combobox = ttk.Combobox(root, values=["一级", "二级", "三级", "四级", "五级", "六级", "七级"])
        self.grade_combobox.grid(row=0, column=3, padx=10, pady=5)

        # 添加下拉选项，输入需要测量的参数
        self.parameter_label = ttk.Label(root, text="请输入需要测量的参数")
        self.parameter_label.grid(row=1, column=0, padx=10, pady=5)
        self.parameter_combobox = ttk.Combobox(root, values=["分度误差", "灵敏度"])
        self.parameter_combobox.grid(row=1, column=1, padx=10, pady=5)
        self.parameter_combobox.bind("<<ComboboxSelected>>", self.extract_measurement_data)

        # 添加按钮，用于触发成本-不确定度分析
        self.calculate_button = ttk.Button(root, text="计算不确定度", command=self.calculate_uncertainty)
        self.calculate_button.grid(row=4, column=0, columnspan=4, padx=10, pady=5)

        # 加载测量参数库Excel文件
        self.measurement_data = self.load_measurement_data()

    def load_measurement_data(self):
        try:
            excel_file_path = "测量参数库.xlsx"
            wb = load_workbook(excel_file_path)
            print("成功加载测量参数库文件")
            measurement_data = {}
            for sheetname in wb.sheetnames:
                ws = wb[sheetname]
                data = {"列标题": [], "内容": []}
                for row in ws.iter_rows(values_only=True):
                    column_titles = row[0]
                    column_content = row[1:]
                    if column_titles is not None:
                        data["列标题"].append(column_titles)
                        data["内容"].append(column_content)
                measurement_data[sheetname] = data
            return measurement_data
        except Exception as e:
            messagebox.showerror("错误", f"无法加载测量参数库文件: {e}")

    def create_word_document(self, event=None):
        try:
            # 检查是否存在"仪器.xlsx"文件
            if not os.path.exists("仪器.xlsx"):
                messagebox.showerror("错误", "找不到仪器.xlsx文件，请先导入该文件")
                return

            # 加载"仪器.xlsx"文件
            wb = load_workbook("仪器.xlsx")
            ws = wb.active

            # 查找所选仪器对应的行
            instrument_name = self.instrument_combobox.get()
            found = False
            for row in ws.iter_rows(values_only=True):
                if row[0] == instrument_name:  # 假设仪器名称在第一列
                    found = True
                    data_to_save = row
                    break

            if not found:
                messagebox.showerror("错误", f"找不到 {instrument_name} 的信息，请检查Excel文件")
                return

            # 读取列标题
            column_titles = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))

            # 创建新的Word文档并保存数据
            doc = Document()

            # 设置标题样式
            title = doc.add_heading(instrument_name, level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.runs[0]
            font = run.font
            font.bold = True  # 设置字体为黑体
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体为黑体

            # 添加仪器信息列标题和内容
            for title, value in zip(column_titles[1:], data_to_save[1:]):
                p = doc.add_paragraph()
                p.add_run(f"{title}：").bold = True  # 设置列标题为粗体
                p.add_run(f"{value}")  # 添加内容
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.space_after = Pt(12)  # 设置段后间距为12磅

            doc.save(f"{instrument_name}.docx")

            messagebox.showinfo("成功", f"已生成以 {instrument_name} 命名的Word文档")

        except Exception as e:
            messagebox.showerror("错误", f"发生错误: {e}")

    def extract_measurement_data(self, event=None):
        parameter_name = self.parameter_combobox.get()
        if not parameter_name:
            messagebox.showerror("错误", "请选择参数名称")
            return
        try:
            instrument_name = self.instrument_combobox.get()
            if not instrument_name:
                messagebox.showerror("错误", "请选择仪器名称")
                return
            if parameter_name == "分度误差":  # 如果选择的参数是"分度误差"
                # 加载测量参数库 Excel 文件
                wb = load_workbook("测量参数库.xlsx")
                ws = wb.active
            else:
                messagebox.showerror("错误", "暂不支持此参数")
                return

            # 查找测量参数对应的行
            found = False
            for row in ws.iter_rows(values_only=True):
                if row[0] == parameter_name:  # 假设参数名称在第一列
                    found = True
                    data_to_save = row
                    break

            if not found:
                messagebox.showerror("错误", f"找不到 {parameter_name} 的信息，请检查Excel文件")
                return

            # 读取列标题
            column_titles = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))

            # 加载仪器信息的 Word 文档
            doc = Document(f"{instrument_name}.docx")

            # 添加测量参数信息列标题和内容
            for title, value in zip(column_titles[1:], data_to_save[1:]):
                p = doc.add_paragraph()
                p.add_run(f"{title}：").bold = True  # 设置列标题为粗体
                p.add_run(f"{value}")  # 添加内容
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.space_after = Pt(12)  # 设置段后间距为12磅

            #保存 Word 文档
            doc.save(f"{instrument_name}.docx")
            messagebox.showinfo("成功", f"已将 {parameter_name} 数据保存到Word文档中")

        except Exception as e:
            messagebox.showerror("错误", f"发生错误: {e}")

    def calculate_uncertainty(self):
        try:
            # 获取所需的参数值
            x = 20  # 测量次数，示例值
            grade = self.grade_combobox.get()
            instrument_name = self.instrument_combobox.get()

            if not instrument_name:
                messagebox.showerror("错误", "请选择仪器名称")
                return

            # 计算不确定度
            uncertainty_a = (0.1 / np.sqrt(x)) ** 2
            uncertainty_b_method = (0.17 / np.sqrt(x)) ** 2

            uncertainty_c_method_one_a = 0.03 ** 2
            uncertainty_c_method_one_b = 0.06 ** 2
            uncertainty_c_method_one_c = 0.29 ** 2

            uncertainty_b_method_two_a = 0.14 ** 2
            uncertainty_b_method_two_b = 0.29 ** 2
            uncertainty_b_method_two_c = 0.57 ** 2

            uncertainty_model_one_a = np.sqrt(uncertainty_a + uncertainty_b_method + uncertainty_c_method_one_a)
            uncertainty_model_one_b = np.sqrt(uncertainty_a + uncertainty_b_method + uncertainty_c_method_one_b)
            uncertainty_model_one_c = np.sqrt(uncertainty_a + uncertainty_b_method + uncertainty_c_method_one_c)

            uncertainty_model_two_a = np.sqrt(uncertainty_a + uncertainty_b_method_two_a)
            uncertainty_model_two_b = np.sqrt(uncertainty_a + uncertainty_b_method_two_b)
            uncertainty_model_two_c = np.sqrt(uncertainty_a + uncertainty_b_method_two_c)

            expanded_uncertainty_model_one_a = 2 * np.sqrt(uncertainty_model_one_a)
            expanded_uncertainty_model_one_b = 2 * np.sqrt(uncertainty_model_one_b)
            expanded_uncertainty_model_one_c = 2 * np.sqrt(uncertainty_model_one_c)

            expanded_uncertainty_model_two_a = 2 * np.sqrt(uncertainty_model_two_a)
            expanded_uncertainty_model_two_b = 2 * np.sqrt(uncertainty_model_two_b)
            expanded_uncertainty_model_two_c = 2 * np.sqrt(uncertainty_model_two_c)

            # 根据所选等级和方法，计算不确定度
            result = f"计算结果：\n\n"
            if grade == "五级" or grade == "六级" or grade == "七级":
                # 方法一
                cost_a = 500 + 500 + 10 * (x - 17)
                cost_b = 500 + 400 + 10 * (x - 16)
                cost_c = 500 + 300 + 10 * (x - 12)

                result += "方法一：\n"
                result += f"仪器A 成本：{cost_a}\n"
                result += f"不确定度：{uncertainty_model_one_a}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_one_a}\n\n"

                result += f"仪器B 成本：{cost_b}\n"
                result += f"不确定度：{uncertainty_model_one_b}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_one_b}\n\n"

                result += f"仪器C 成本：{cost_c}\n"
                result += f"不确定度：{uncertainty_model_one_c}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_one_c}\n\n"
                # 方法二
                cost_a = 400 + 400 + 10 * (x - 17)
                cost_b = 300 + 300 + 10 * (x - 16)
                cost_c = 200 + 200 + 10 * (x - 12)

                result += "方法二：\n"
                result += f"仪器A 成本：{cost_a}\n"
                result += f"不确定度：{uncertainty_model_two_a}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_two_a}\n\n"

                result += f"仪器B 成本：{cost_b}\n"
                result += f"不确定度：{uncertainty_model_two_b}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_two_b}\n\n"

                result += f"仪器C 成本：{cost_c}\n"
                result += f"不确定度：{uncertainty_model_two_c}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_two_c}\n\n"
            elif grade == "四级":
                # 方法一
                cost_a = 500 + 500 + 10 * (x - 17)
                cost_b = 500 + 400 + 10 * (x - 16)

                result += "方法一：\n"
                result += f"仪器A 成本：{cost_a}\n"
                result += f"不确定度：{uncertainty_model_one_a}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_one_a}\n\n"

                result += f"仪器B 成本：{cost_b}\n"
                result += f"不确定度：{uncertainty_model_one_b}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_one_b}\n\n"

                # 方法二
                cost_a = 400 + 400 + 10 * (x - 17)
                cost_b = 300 + 300 + 10 * (x - 16)

                result += "方法二：\n"
                result += f"仪器A 成本：{cost_a}\n"
                result += f"不确定度：{uncertainty_model_two_a}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_two_a}\n\n"

                result += f"仪器B 成本：{cost_b}\n"
                result += f"不确定度：{uncertainty_model_two_b}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_two_b}\n\n"

                pass
            elif grade in ["一级", "二级", "三级"]:
                # 方法一
                cost_a = 500 + 500 + 10 * (x - 17)

                result += "方法一：\n"
                result += f"仪器A 成本：{cost_a}\n"
                result += f"不确定度：{uncertainty_model_one_a}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_one_a}\n\n"

                # 方法二
                cost_a = 400 + 400 + 10 * (x - 17)

                result += "方法二：\n"
                result += f"仪器A 成本：{cost_a}\n"
                result += f"不确定度：{uncertainty_model_two_a}\n"
                result += f"扩展不确定度：{expanded_uncertainty_model_two_a}\n\n"

            # 将结果添加到之前创建的 Word 文档
            doc = Document(f"{instrument_name}.docx")

            # 加载仪器信息的 Word 文档
            instrument_doc = Document(f"{instrument_name}.docx")

            # 加载测量参数的 Word 文档
            parameter_name = self.parameter_combobox.get()
            parameter_doc = Document(f"{parameter_name}.docx")

            # 将测量参数文档的内容添加到仪器信息文档中
            for element in parameter_doc.element.body:
                instrument_doc.element.body.append(element)

            # 保存更新后的 Word 文档
            instrument_doc.save(f"{instrument_name}.docx")

            messagebox.showinfo("成功", "不确定度计算完成")

            p = doc.add_paragraph()
            run = p.add_run(result)
            font = run.font
            font.size = Pt(12)  # 设置字体大小为12磅
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体为宋体
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_after = Pt(12)  # 设置段后间距为12磅

            doc.save(f"{instrument_name}.docx")

            messagebox.showinfo("成功", "不确定度计算完成")

        except Exception as e:
            messagebox.showerror("错误", f"发生错误: {e}")


def main():
    root = tk.Tk()
    app = DynamicProcedureApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
